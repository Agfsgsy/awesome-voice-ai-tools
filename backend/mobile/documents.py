"""استخراج نصوص المستندات وتهيئة الأرقام والتواريخ والعملات للنطق العربي."""

from __future__ import annotations

import re
from collections.abc import Iterable
from datetime import datetime, timezone
from pathlib import Path

from backend.mobile.config import SUPPORTED_DOCUMENT_EXTENSIONS


class DocumentReadError(ValueError):
    """مستند غير صالح أو غير مدعوم."""


ARABIC_MONTHS = {
    1: "يناير",
    2: "فبراير",
    3: "مارس",
    4: "أبريل",
    5: "مايو",
    6: "يونيو",
    7: "يوليو",
    8: "أغسطس",
    9: "سبتمبر",
    10: "أكتوبر",
    11: "نوفمبر",
    12: "ديسمبر",
}

CURRENCY_NAMES = {
    "$": "دولار أمريكي",
    "USD": "دولار أمريكي",
    "SAR": "ريال سعودي",
    "YER": "ريال يمني",
    "AED": "درهم إماراتي",
    "EUR": "يورو",
    "GBP": "جنيه إسترليني",
    "KWD": "دينار كويتي",
    "QAR": "ريال قطري",
}

ONES = [
    "صفر",
    "واحد",
    "اثنان",
    "ثلاثة",
    "أربعة",
    "خمسة",
    "ستة",
    "سبعة",
    "ثمانية",
    "تسعة",
    "عشرة",
    "أحد عشر",
    "اثنا عشر",
    "ثلاثة عشر",
    "أربعة عشر",
    "خمسة عشر",
    "ستة عشر",
    "سبعة عشر",
    "ثمانية عشر",
    "تسعة عشر",
]
TENS = {20: "عشرون", 30: "ثلاثون", 40: "أربعون", 50: "خمسون", 60: "ستون", 70: "سبعون", 80: "ثمانون", 90: "تسعون"}
HUNDREDS = {
    100: "مئة",
    200: "مئتان",
    300: "ثلاثمئة",
    400: "أربعمئة",
    500: "خمسمئة",
    600: "ستمئة",
    700: "سبعمئة",
    800: "ثمانمئة",
    900: "تسعمئة",
}
SCALES = [
    ("", "", ""),
    ("ألف", "ألفان", "آلاف"),
    ("مليون", "مليونان", "ملايين"),
    ("مليار", "ملياران", "مليارات"),
    ("تريليون", "تريليونان", "تريليونات"),
]


def _join(parts: Iterable[str]) -> str:
    return " و".join(part for part in parts if part)


def _under_thousand(value: int) -> str:
    if value < 20:
        return ONES[value]
    parts: list[str] = []
    hundreds = value // 100 * 100
    remainder = value % 100
    if hundreds:
        parts.append(HUNDREDS[hundreds])
    if remainder:
        if remainder < 20:
            parts.append(ONES[remainder])
        else:
            unit = remainder % 10
            if unit:
                parts.append(ONES[unit])
            parts.append(TENS[remainder // 10 * 10])
    return _join(parts)


def number_to_arabic_words(value: int) -> str:
    if value == 0:
        return ONES[0]
    if value < 0:
        return "سالب " + number_to_arabic_words(abs(value))
    if value >= 10**15:
        return str(value)
    groups: list[int] = []
    remaining = value
    while remaining:
        groups.append(remaining % 1000)
        remaining //= 1000
    parts: list[str] = []
    for index in range(len(groups) - 1, -1, -1):
        group = groups[index]
        if not group:
            continue
        if index == 0:
            parts.append(_under_thousand(group))
            continue
        singular, dual, plural = SCALES[index]
        if group == 1:
            parts.append(singular)
        elif group == 2:
            parts.append(dual)
        elif 3 <= group <= 10:
            parts.append(f"{_under_thousand(group)} {plural}")
        else:
            parts.append(f"{_under_thousand(group)} {singular}")
    return _join(parts)


def _decimal_to_words(raw: str) -> str:
    cleaned = raw.replace(",", "")
    if "." not in cleaned:
        return number_to_arabic_words(int(cleaned))
    whole, fraction = cleaned.split(".", 1)
    fraction_words = " ".join(ONES[int(digit)] for digit in fraction if digit.isdigit())
    return f"{number_to_arabic_words(int(whole or '0'))} فاصلة {fraction_words}"


def normalize_text_for_speech(text: str) -> str:
    def currency_replace(match: re.Match[str]) -> str:
        prefix, amount, suffix = match.group(1), match.group(2), match.group(3)
        currency = prefix or suffix
        return f"{_decimal_to_words(amount)} {CURRENCY_NAMES.get(currency.upper(), currency)}"

    currency_pattern = re.compile(
        r"(?:(\$|USD|SAR|YER|AED|EUR|GBP|KWD|QAR)\s*)?([0-9][0-9,]*(?:\.[0-9]+)?)\s*(USD|SAR|YER|AED|EUR|GBP|KWD|QAR)?",
        re.IGNORECASE,
    )

    def guarded_currency(match: re.Match[str]) -> str:
        if not match.group(1) and not match.group(3):
            return match.group(0)
        return currency_replace(match)

    value = currency_pattern.sub(guarded_currency, text)

    def date_replace(match: re.Match[str]) -> str:
        raw = match.group(0)
        formats = ("%Y-%m-%d", "%d/%m/%Y", "%d-%m-%Y")
        parsed = next(
            (datetime.strptime(raw, fmt).replace(tzinfo=timezone.utc) for fmt in formats if _can_parse(raw, fmt)),
            None,
        )
        if parsed is None:
            return raw
        return (
            f"{number_to_arabic_words(parsed.day)} {ARABIC_MONTHS[parsed.month]} {number_to_arabic_words(parsed.year)}"
        )

    value = re.sub(r"\b(?:\d{4}-\d{1,2}-\d{1,2}|\d{1,2}[/-]\d{1,2}[/-]\d{4})\b", date_replace, value)

    def number_replace(match: re.Match[str]) -> str:
        try:
            return _decimal_to_words(match.group(0))
        except (ValueError, OverflowError):
            return match.group(0)

    return re.sub(r"\b\d+(?:\.\d+)?\b", number_replace, value)


def _can_parse(value: str, fmt: str) -> bool:
    try:
        datetime.strptime(value, fmt).replace(tzinfo=timezone.utc)
        return True
    except ValueError:
        return False


def extract_document(path: Path, max_characters: int = 120_000) -> str:
    extension = path.suffix.lower()
    if extension not in SUPPORTED_DOCUMENT_EXTENSIONS:
        raise DocumentReadError("صيغة المستند غير مدعومة؛ استخدم PDF أو DOCX أو TXT")
    try:
        if extension == ".txt":
            text = _read_text_file(path)
        elif extension == ".pdf":
            from pypdf import PdfReader

            reader = PdfReader(str(path))
            text = "\n\n".join((page.extract_text() or "") for page in reader.pages)
        else:
            from docx import Document

            document = Document(str(path))
            paragraphs = [paragraph.text for paragraph in document.paragraphs]
            for table in document.tables:
                for row in table.rows:
                    paragraphs.append(" | ".join(cell.text for cell in row.cells))
            text = "\n".join(paragraphs)
    except ImportError as exc:
        raise DocumentReadError("اعتمادات قراءة المستندات غير مثبتة على الخادم") from exc
    except Exception as exc:
        raise DocumentReadError("تعذر قراءة المستند أو أنه تالف") from exc
    normalized = re.sub(r"[ \t]+", " ", text).strip()
    if not normalized:
        raise DocumentReadError("المستند لا يحتوي نصًا قابلاً للقراءة")
    if len(normalized) > max_characters:
        raise DocumentReadError(f"المستند أطول من الحد المسموح ({max_characters} حرف)")
    return normalized


def _read_text_file(path: Path) -> str:
    raw = path.read_bytes()
    for encoding in ("utf-8-sig", "utf-16", "cp1256", "iso-8859-6"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    raise DocumentReadError("ترميز ملف TXT غير مدعوم")


def split_text(text: str, limit: int = 3500) -> list[str]:
    if limit < 200:
        raise ValueError("حد المقطع النصي صغير جدًا")
    paragraphs = [item.strip() for item in re.split(r"\n+", text) if item.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        sentences = re.split(r"(?<=[.!؟؛])\s+", paragraph)
        for sentence in sentences:
            if len(sentence) > limit:
                words = sentence.split()
                for word in words:
                    candidate = f"{current} {word}".strip()
                    if len(candidate) > limit and current:
                        chunks.append(current)
                        current = word
                    else:
                        current = candidate
                continue
            candidate = f"{current}\n{sentence}".strip()
            if len(candidate) > limit and current:
                chunks.append(current)
                current = sentence
            else:
                current = candidate
    if current:
        chunks.append(current)
    return chunks
